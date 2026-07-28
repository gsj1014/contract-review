# -*- coding: utf-8 -*-
"""
专业版 Word 报告生成器 - Phase 3+
功能：生成精美格式的合同审核报告 Word 文档
"""

import os
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ProfessionalReportGenerator:
    """
    专业版 Word 报告生成器
    
    生成符合法律文书规范的高质量审核报告：
    - 仿宋字体，1.5倍行距
    - 清晰的章节结构
    - 专业的表格样式
    - 彩色风险等级标识
    """
    
    # 颜色定义
    COLORS = {
        'high_risk': RGBColor(192, 0, 0),      # 红色
        'medium_risk': RGBColor(255, 192, 0),   # 橙色
        'low_risk': RGBColor(0, 176, 80),       # 绿色
        'info': RGBColor(0, 112, 192),          # 蓝色
        'title': RGBColor(0, 0, 0),              # 黑色
        'subtitle': RGBColor(64, 64, 64),        # 深灰色
        'legal': RGBColor(112, 48, 160),        # 紫色（法律依据）
    }
    
    # 字体定义
    FONTS = {
        'title': '黑体',
        'heading': '黑体',
        'body': '仿宋',
        'emphasis': '楷体',
    }

    def __init__(self, review_data: Dict, title: str = "合同审核报告"):
        """
        初始化报告生成器
        
        Args:
            review_data: 审核结果数据
            title: 报告标题
        """
        self.review_data = review_data
        self.title = title
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档基本格式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = self.FONTS['body']
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.FONTS['body'])
        
        # 设置段落间距
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    def _set_run_font(self, run, font_name: str, size: Pt = None, bold: bool = False,
                      color: RGBColor = None, italic: bool = False):
        """设置文本运行格式"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        if size:
            run.font.size = size
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if color:
            run.font.color.rgb = color
    
    def _add_title(self, text: str, level: int = 1):
        """添加标题"""
        heading = self.doc.add_heading('', level=level)
        run = heading.add_run(text)
        
        if level == 0:
            self._set_run_font(run, self.FONTS['title'], Pt(22), bold=True)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 1:
            self._set_run_font(run, self.FONTS['heading'], Pt(16), bold=True)
        elif level == 2:
            self._set_run_font(run, self.FONTS['heading'], Pt(14), bold=True)
        else:
            self._set_run_font(run, self.FONTS['body'], Pt(12), bold=True)
        
        return heading
    
    def _add_paragraph(self, text: str = '', bold: bool = False, 
                       italic: bool = False, alignment: WD_ALIGN_PARAGRAPH = None,
                       space_before: Pt = None, space_after: Pt = None):
        """添加段落"""
        para = self.doc.add_paragraph()
        
        if text:
            run = para.add_run(text)
            self._set_run_font(run, self.FONTS['body'], Pt(12), bold=bold, italic=italic)
        
        if alignment:
            para.alignment = alignment
        
        if space_before is not None:
            para.paragraph_format.space_before = space_before
        if space_after is not None:
            para.paragraph_format.space_after = space_after
        
        return para
    
    def _add_styled_paragraph(self, text: str, style_name: str = 'BodyText'):
        """添加带样式的段落"""
        para = self.doc.add_paragraph(style=style_name)
        run = para.add_run(text)
        self._set_run_font(run, self.FONTS['body'], Pt(12))
        return para
    
    def _create_risk_table(self, risks: List[Dict], title: str = None) -> None:
        """创建风险表格"""
        if title:
            self._add_title(title, level=2)
        
        if not risks:
            self._add_paragraph('未发现相关风险。', italic=True)
            return
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置列宽
        for i, width in enumerate([Cm(1.5), Cm(3), Cm(4.5), Cm(6)]):
            table.columns[i].width = width
        
        # 表头
        header_cells = table.rows[0].cells
        headers = ['风险等级', '条款类型', '问题描述', '修改建议']
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.bold = True
            self._set_run_font(run, self.FONTS['body'], Pt(10.5))
            
            # 表头背景色
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9E2F3')
            cell._tc.get_or_add_tcPr().append(shading)
        
        # 数据行
        for risk in risks:
            row_cells = table.add_row().cells
            
            # 风险等级（带颜色）
            level_text = risk.get('level', '')
            row_cells[0].text = level_text
            self._color_risk_level(row_cells[0].paragraphs[0].runs[0], level_text)
            
            # 条款类型
            row_cells[1].text = risk.get('type', '')
            self._style_table_cell(row_cells[1], Pt(10.5))
            
            # 问题描述
            row_cells[2].text = risk.get('description', '')
            self._style_table_cell(row_cells[2], Pt(10.5))
            
            # 修改建议
            row_cells[3].text = risk.get('suggestion', '')
            self._style_table_cell(row_cells[3], Pt(10.5))
    
    def _color_risk_level(self, run, level_text: str):
        """根据风险等级设置颜色"""
        self._set_run_font(run, self.FONTS['body'], Pt(10.5), bold=True)
        
        if '🔴' in level_text or '高风险' in level_text:
            run.font.color.rgb = self.COLORS['high_risk']
        elif '🟡' in level_text or '中风险' in level_text:
            run.font.color.rgb = self.COLORS['medium_risk']
        elif '🟢' in level_text or '低风险' in level_text:
            run.font.color.rgb = self.COLORS['low_risk']
        else:
            run.font.color.rgb = self.COLORS['info']
    
    def _style_table_cell(self, cell, size: Pt = Pt(10.5)):
        """设置表格单元格样式"""
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        if para.runs:
            run = para.runs[0]
            self._set_run_font(run, self.FONTS['body'], size)
    
    def _add_gap_table(self, gaps: List[Dict]) -> None:
        """创建差异比对表格"""
        self._add_title('五、📋 模板比对分析', level=2)
        
        if not gaps:
            self._add_paragraph('未发现明显差异。', italic=True)
            return
        
        # 按类型分组
        missing = [g for g in gaps if g.get('type') == '缺失']
        different = [g for g in gaps if g.get('type') == '偏离']
        
        # 缺失条款
        if missing:
            self._add_title('⚠️ 缺失条款', level=3)
            self._create_gap_table(missing)
        
        # 偏离条款
        if different:
            self._add_title('⚡ 偏离条款', level=3)
            self._create_gap_table(different)
    
    def _create_gap_table(self, gaps: List[Dict]) -> None:
        """创建差异表格"""
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['条款名称', '重要性', '差异说明', '法律依据']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.bold = True
            self._set_run_font(run, self.FONTS['body'], Pt(10.5))
            
            # 表头背景色
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'E2EFDA')
            cell._tc.get_or_add_tcPr().append(shading)
        
        # 数据行
        for gap in gaps:
            row_cells = table.add_row().cells
            row_cells[0].text = gap.get('clause', '')
            row_cells[1].text = gap.get('importance', '')
            row_cells[2].text = gap.get('difference', '')
            row_cells[3].text = gap.get('legal_basis', '')
            
            for cell in row_cells:
                self._style_table_cell(cell, Pt(10.5))
    
    def _add_legal_compliance(self, legal_results: List[Dict]) -> None:
        """添加法律合规审核章节"""
        if not legal_results:
            return
        
        self._add_title('六、⚖️ 法律合规审核报告', level=2)
        
        # 按来源分组
        by_source = {}
        for item in legal_results:
            source = item.get('legal_source', '行业规范')
            by_source.setdefault(source, []).append(item)
        
        # 统计摘要
        self._add_paragraph('法律合规检查统计：')
        
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        headers = ['法律来源', '问题类型', '数量']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.bold = True
            self._set_run_font(run, self.FONTS['body'], Pt(10.5))
        
        for source, items in by_source.items():
            row = table.add_row().cells
            row[0].text = source
            row[1].text = '缺失/不合规'
            row[2].text = str(len(items))
            for cell in row:
                self._style_table_cell(cell)
        
        # 详细问题列表
        for source, items in by_source.items():
            self._add_paragraph(f'\n【{source}】', bold=True)
            for item in items:
                status_icon = '❌ ' if item.get('status') == 'non_compliant' else '⚠️ '
                self._add_paragraph(f'{status_icon}{item.get("clause", "")}', bold=True)
                self._add_paragraph(f'  法律依据：{item.get("legal_basis", "N/A")}')
                self._add_paragraph(f'  审核发现：{item.get("findings", "N/A")}')
                self._add_paragraph(f'  合规建议：{item.get("suggestion", "N/A")}')
    
    def _add_summary_table(self, summary: Dict) -> None:
        """添加摘要统计表"""
        self._add_title('审核结论', level=1)
        
        # 评估结论
        self._add_paragraph(summary.get('overall_assessment', '无法生成结论'))
        self._add_paragraph('')
        
        # 统计数据
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        data = [
            ('🔴 高风险', f'{summary.get("high_risk_count", 0)} 项'),
            ('🟡 中风险', f'{summary.get("medium_risk_count", 0)} 项'),
            ('🟢 低风险', f'{summary.get("low_risk_count", 0)} 项'),
            ('📊 风险评分', f'{self.review_data.get("risk_score", 0):.0f}/100'),
        ]
        
        for i, (label, value) in enumerate(data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # 设置样式
            for cell in row.cells:
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0]
                self._set_run_font(run, self.FONTS['body'], Pt(12))
    
    def _add_basic_info(self, basic_info: Dict) -> None:
        """添加合同基本信息"""
        self._add_title('二、合同基本信息', level=1)
        
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        fields = [
            ('合同名称', basic_info.get('contract_name', '未明确')),
            ('甲方（委托方）', basic_info.get('party_a', '未明确')),
            ('乙方（受托方）', basic_info.get('party_b', '未明确')),
            ('签署日期', basic_info.get('signing_date', '未明确')),
            ('合同金额', basic_info.get('contract_amount', '未明确')),
        ]
        
        for i, (field, value) in enumerate(fields):
            row = table.rows[i]
            row.cells[0].text = field
            row.cells[1].text = value
            
            # 标题列样式
            row.cells[0].paragraphs[0].runs[0].bold = True
            self._style_table_cell(row.cells[0])
            self._style_table_cell(row.cells[1])
            
            # 标题列背景色
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F2F2F2')
            row.cells[0]._tc.get_or_add_tcPr().append(shading)
    
    def generate(self, output_path: str = None) -> str:
        """
        生成完整报告
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            生成的报告文件路径
        """
        # 文档标题
        self._add_title(self.title, level=0)
        
        # 生成时间
        time_para = self.doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        time_run = time_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
        self._set_run_font(time_run, self.FONTS['body'], Pt(10), italic=True)
        
        # 元信息
        meta_para = self.doc.add_paragraph()
        meta_run = meta_para.add_run(
            f"合同类型：{self.review_data.get('classified', {}).get('contract_type', '未识别')}  |  "
            f"风险评分：{self.review_data.get('risk_score', 0):.0f}/100"
        )
        self._set_run_font(meta_run, self.FONTS['body'], Pt(11))
        
        self.doc.add_paragraph()  # 空行
        
        # 审核结论
        summary = self.review_data.get('reviewed', {}).get('summary', {})
        self._add_summary_table(summary)
        
        # 合同基本信息
        basic_info = self.review_data.get('parsed', {}).get('basic_info', {})
        self._add_basic_info(basic_info)
        
        # 风险条款
        risks = self.review_data.get('reviewed', {}).get('risks', {})
        
        high_risks = risks.get('HIGH', [])
        if high_risks:
            self._add_title('三、🔴 高风险条款（必须修改）', level=1)
            self._create_risk_table(high_risks)
        
        medium_risks = risks.get('MEDIUM', [])
        if medium_risks:
            self._add_title('四、🟡 中风险条款（建议修改）', level=1)
            self._create_risk_table(medium_risks)
        
        # 差异比对
        if self.review_data.get('gap_result'):
            gaps = self.review_data['gap_result'].get('gaps', [])
            self._add_gap_table(gaps)
        
        # 法律合规
        if self.review_data.get('legal_results'):
            self._add_legal_compliance(self.review_data['legal_results'])
        
        # 低风险
        low_risks = risks.get('LOW', [])
        if low_risks:
            self._add_title('七、🟢 低风险条款（可保留参考）', level=1)
            self._create_risk_table(low_risks)
        
        # 页脚
        self.doc.add_paragraph()
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(
            '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
            '本报告由 WorkBuddy 合同审核 Skill 自动生成\n'
            '仅供参考，最终意见以执业律师判断为准'
        )
        self._set_run_font(footer_run, self.FONTS['body'], Pt(9), 
                          color=RGBColor(128, 128, 128), italic=True)
        
        # 保存
        if not output_path:
            base_dir = os.path.dirname(os.path.dirname(__file__))
            output_dir = os.path.join(base_dir, 'outputs')
            os.makedirs(output_dir, exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = os.path.join(output_dir, f'审核报告_{timestamp}.docx')
        
        self.doc.save(output_path)
        return output_path


def create_professional_report(review_data: Dict, output_path: str = None) -> str:
    """
    创建专业版 Word 审核报告
    
    Args:
        review_data: 审核结果数据
        output_path: 输出路径
        
    Returns:
        报告文件路径
    """
    generator = ProfessionalReportGenerator(review_data)
    return generator.generate(output_path)
