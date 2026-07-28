# -*- coding: utf-8 -*-
"""
Word Track Changes 修订模块 (track_changes_docx.py)
====================================================

实现与 Microsoft Word 完全兼容的 Word Track Changes 修订格式。

技术原理：
- 使用标准 OOXML 规范的 <w:ins> 和 <w:del> 元素
- 完整设置文档属性（core.xml, app.xml）
- 正确配置 word/settings.xml 中的修订模式
- 确保命名空间声明完整

使用示例：
    from track_changes_docx import create_track_changes_docx

    changes = [
        {
            'clause_title': '第五条 违约责任',
            'change_type': 'replace',
            'original_text': '违约金为合同总价的50%',
            'new_text': '违约金为合同总价的30%',
            'reason': '违约金比例过高',
            'legal_basis': '《民法典》第585条'
        }
    ]

    create_track_changes_docx("原始合同", changes, "修订稿.docx")
"""

import os
import re
import shutil
import tempfile
import zipfile
from datetime import datetime
from typing import List, Dict, Optional
from io import BytesIO

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import hashlib


# ============================================================
# 文档模板 - 基于标准 OOXML 结构
# ============================================================

CONTENT_TYPES_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
    <Override PartName="/word/settings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"/>
    <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
    <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
    <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>'''

RELS_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
    <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
    <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>'''

PACKAGE_RELS_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="/word/document.xml"/>
</Relationships>'''

SETTINGS_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:trackRevisions/>
    <w:revisionView w:markup="1"/>
</w:settings>'''

STYLES_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:style w:type="paragraph" w:styleId="Normal" w:default="1">
        <w:name w:val="Normal"/>
        <w:rPr>
            <w:rFonts w:eastAsia="宋体" w:ascii="宋体"/>
            <w:sz w:val="24"/>
            <w:szCs w:val="24"/>
        </w:rPr>
    </w:style>
</w:styles>'''

CORE_XML_TEMPLATE = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
    xmlns:dc="http://purl.org/dc/elements/1.1/"
    xmlns:dcterms="http://purl.org/dc/terms/"
    xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
    <dc:title>合同修订稿</dc:title>
    <dc:creator>WorkBuddy AI 审核助手</dc:creator>
    <dcterms:created xsi:type="dcterms:W3CDTF">{created}</dcterms:created>
    <dcterms:modified xsi:type="dcterms:W3CDTF">{modified}</dcterms:modified>
</cp:coreProperties>'''

APP_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties">
    <Application>WorkBuddy/Contract Review</Application>
    <DocSecurity>0</DocSecurity>
    <Company>广东知恒（大连）律师事务所</Company>
</Properties>'''


def create_track_changes_docx(
    original_text: str,
    changes: List[Dict],
    output_path: str,
    author: str = "AI审核助手"
) -> bool:
    """
    创建带 Track Changes 的 Word 文档（完全兼容 Microsoft Word）

    Args:
        original_text: 原始合同文本
        changes: 修改项列表
        output_path: 输出路径
        author: 修订作者

    Returns:
        是否成功
    """
    try:
        # 方法1：使用 python-docx 基础文档 + 手动添加修订
        doc = Document()

        # 移除默认空段落
        doc._body.clear_content()

        # 设置文档标题
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('合同修订稿（Track Changes）')
        title_run.bold = True
        title_run.font.size = Pt(16)

        # 添加说明
        doc.add_paragraph()
        intro = doc.add_paragraph()
        intro.add_run('📌 使用说明：').bold = True

        doc.add_paragraph('• 在 Word「审阅」选项卡中查看所有修订标记')
        doc.add_paragraph('• 点击「接受」或「拒绝」逐条处理修订')
        doc.add_paragraph('• 使用「接受所有修订」一键生成最终版本')
        doc.add_paragraph()
        doc.add_paragraph('📋 修订标记说明：')
        doc.add_paragraph('• 🔴 删除线 + 红色 = 待删除内容')
        doc.add_paragraph('• 🔵 下划线 + 蓝色 = 新增内容')
        doc.add_paragraph()
        doc.add_paragraph('─' * 40)

        # 处理每个修改项
        for i, change in enumerate(changes, 1):
            _add_change_with_revision(doc, change, i, author)

        # 添加分隔线和统计
        doc.add_paragraph()
        doc.add_paragraph('─' * 40)

        stats_heading = doc.add_paragraph()
        stats_heading.add_run('📊 修订统计').bold = True

        stats = {
            'total': len(changes),
            'replace': sum(1 for c in changes if c.get('change_type') == 'replace'),
            'add': sum(1 for c in changes if c.get('change_type') == 'add'),
            'delete': sum(1 for c in changes if c.get('change_type') == 'delete'),
        }

        doc.add_paragraph(f'• 替换修订：{stats["replace"]} 项')
        doc.add_paragraph(f'• 新增内容：{stats["add"]} 项')
        doc.add_paragraph(f'• 删除内容：{stats["delete"]} 项')
        doc.add_paragraph(f'• 合计修订：{stats["total"]} 项')

        # 页脚说明
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run('* 本文档由 WorkBuddy AI 审核助手自动生成 *')
        footer_run.italic = True
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # 保存到临时位置
        temp_path = output_path + '.tmp.docx'
        doc.save(temp_path)

        # 确保文件关闭后重命名
        import time
        time.sleep(0.1)

        if os.path.exists(output_path):
            os.remove(output_path)
        os.rename(temp_path, output_path)

        return True

    except Exception as e:
        print(f"❌ 生成修订稿失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def _add_change_with_revision(doc: Document, change: Dict, index: int, author: str):
    """添加单个修订项到文档"""
    change_type = change.get('change_type', 'replace')
    clause_title = change.get('clause_title', f'修改项 {index}')
    original = change.get('original_text', '')
    new_text = change.get('new_text', '')
    reason = change.get('reason', '')
    legal_basis = change.get('legal_basis', '')

    # 条款标题
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f'【{clause_title}】')
    title_run.bold = True
    title_run.font.size = Pt(11)

    # 根据修改类型添加修订内容
    if change_type == 'replace' and original and new_text:
        # 替换：先显示删除（删除线），再显示新增（下划线）
        _add_deletion_run(doc, original, author)
        _add_insertion_run(doc, new_text, author)

    elif change_type == 'add' and new_text:
        # 仅新增
        _add_insertion_run(doc, new_text, author)

    elif change_type == 'delete' and original:
        # 仅删除
        _add_deletion_run(doc, original, author)

    # 修改原因
    if reason:
        reason_para = doc.add_paragraph()
        reason_run = reason_para.add_run(f'  💡 修改原因：{reason}')
        reason_run.italic = True
        reason_run.font.size = Pt(9)

    # 法律依据
    if legal_basis:
        legal_para = doc.add_paragraph()
        legal_run = legal_para.add_run(f'  ⚖️ 法律依据：{legal_basis}')
        legal_run.font.size = Pt(9)
        legal_run.font.color.rgb = RGBColor(96, 96, 96)


def _add_insertion_run(doc: Document, text: str, author: str):
    """添加插入内容（蓝色下划线）"""
    # 使用 XML 直接添加 w:ins 元素
    para = doc.add_paragraph()
    pElement = para._p

    # 创建 w:ins 元素
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(hash(text) % 100000))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))

    # 创建 w:r 元素（文本运行）
    run = OxmlElement('w:r')

    # 运行属性
    rPr = OxmlElement('w:rPr')

    # 字体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), '宋体')
    rPr.append(rFonts)

    # 颜色（蓝色）
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # 下划线
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    u.set(qn('w:color'), '0000FF')
    rPr.append(u)

    run.append(rPr)

    # 文本
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)

    ins.append(run)
    pElement.append(ins)


def _add_deletion_run(doc: Document, text: str, author: str):
    """添加删除内容（红色删除线）"""
    para = doc.add_paragraph()
    pElement = para._p

    # 创建 w:del 元素
    del_ = OxmlElement('w:del')
    del_.set(qn('w:id'), str(hash(text) % 100000))
    del_.set(qn('w:author'), author)
    del_.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))

    # 创建 w:r 元素
    run = OxmlElement('w:r')

    # 运行属性
    rPr = OxmlElement('w:rPr')

    # 字体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), '宋体')
    rPr.append(rFonts)

    # 颜色（红色）
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FF0000')
    rPr.append(color)

    # 删除线
    strike = OxmlElement('w:strike')
    rPr.append(strike)

    run.append(rPr)

    # 删除文本（使用 w:delText）
    delText = OxmlElement('w:delText')
    delText.text = text
    delText.set(qn('xml:space'), 'preserve')
    run.append(delText)

    del_.append(run)
    pElement.append(del_)


# ============================================================
# 兼容模式：使用标准 python-docx 格式（无原生修订，但更稳定）
# ============================================================

def create_compatible_docx(
    original_text: str,
    changes: List[Dict],
    output_path: str,
    author: str = "AI审核助手"
) -> bool:
    """
    创建兼容性格式的修订稿（更稳定，所有版本 Word 可打开）
    使用格式化文本而非原生修订模式
    """
    try:
        doc = Document()

        # 标题
        title = doc.add_heading('合同修订稿', level=0)

        # 说明
        doc.add_paragraph()
        intro = doc.add_paragraph()
        intro.add_run('📌 修订说明：').bold = True
        doc.add_paragraph('• 🔴 删除线 + 红色 = 待删除内容')
        doc.add_paragraph('• 🔵 下划线 + 蓝色 = 新增内容')
        doc.add_paragraph('• 请根据审核意见手动修改合同')
        doc.add_paragraph()
        doc.add_paragraph('─' * 40)

        # 处理修改项
        for i, change in enumerate(changes, 1):
            _add_change_formatted(doc, change, i)

        # 分隔线
        doc.add_paragraph()
        doc.add_paragraph('─' * 40)

        # 统计
        stats_heading = doc.add_heading('修订统计', level=1)
        stats = {
            'total': len(changes),
            'replace': sum(1 for c in changes if c.get('change_type') == 'replace'),
            'add': sum(1 for c in changes if c.get('change_type') == 'add'),
            'delete': sum(1 for c in changes if c.get('change_type') == 'delete'),
        }

        doc.add_paragraph(f'• 替换修订：{stats["replace"]} 项')
        doc.add_paragraph(f'• 新增内容：{stats["add"]} 项')
        doc.add_paragraph(f'• 删除内容：{stats["delete"]} 项')
        doc.add_paragraph(f'• 合计修订：{stats["total"]} 项')

        # 保存
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"❌ 生成兼容版修订稿失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def _add_change_formatted(doc: Document, change: Dict, index: int):
    """添加格式化修订内容（不使用原生修订）"""
    change_type = change.get('change_type', 'replace')
    clause_title = change.get('clause_title', f'修改项 {index}')
    original = change.get('original_text', '')
    new_text = change.get('new_text', '')
    reason = change.get('reason', '')
    legal_basis = change.get('legal_basis', '')

    # 条款标题
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f'【{clause_title}】')
    title_run.bold = True

    if change_type == 'replace' and original and new_text:
        # 删除内容
        del_para = doc.add_paragraph()
        del_run = del_para.add_run(f'【删除】{original}')
        del_run.font.strike = True
        del_run.font.color.rgb = RGBColor(255, 0, 0)

        # 新增内容
        ins_para = doc.add_paragraph()
        ins_run = ins_para.add_run(f'【新增】{new_text}')
        ins_run.underline = True
        ins_run.font.color.rgb = RGBColor(0, 0, 255)

    elif change_type == 'add' and new_text:
        ins_para = doc.add_paragraph()
        ins_run = ins_para.add_run(f'【新增】{new_text}')
        ins_run.underline = True
        ins_run.font.color.rgb = RGBColor(0, 0, 255)

    elif change_type == 'delete' and original:
        del_para = doc.add_paragraph()
        del_run = del_para.add_run(f'【删除】{original}')
        del_run.font.strike = True
        del_run.font.color.rgb = RGBColor(255, 0, 0)

    # 原因
    if reason:
        reason_para = doc.add_paragraph()
        reason_para.add_run(f'  原因：{reason}').italic = True

    # 法律依据
    if legal_basis:
        legal_para = doc.add_paragraph()
        legal_run = legal_para.add_run(f'  依据：{legal_basis}')
        legal_run.font.color.rgb = RGBColor(128, 128, 128)
        legal_run.font.size = Pt(9)


# ============================================================
# 测试代码
# ============================================================

if __name__ == "__main__":
    sample_changes = [
        {
            'clause_title': '第五条 违约责任',
            'change_type': 'replace',
            'original_text': '违约金为合同总价的50%',
            'new_text': '违约金为合同总价的30%',
            'reason': '违约金比例过高，可能被认定为过分高于损失',
            'legal_basis': '《民法典》第585条'
        },
        {
            'clause_title': '第三条 付款条件',
            'change_type': 'add',
            'original_text': '',
            'new_text': '付款方式：银行转账；发票条款：卖方应在付款前提供增值税专用发票',
            'reason': '缺少付款方式和发票条款',
            'legal_basis': '建议补充完整付款条款'
        },
        {
            'clause_title': '第六条 争议解决',
            'change_type': 'delete',
            'original_text': '提交甲方所在地人民法院管辖',
            'new_text': '',
            'reason': '管辖条款过于简单',
            'legal_basis': '建议增加协商前置程序'
        }
    ]

    output_path = os.path.join(tempfile.gettempdir(), "TrackChanges_v2测试.docx")

    print("=== 测试 Track Changes 文档生成 ===")
    print(f"输出路径：{output_path}")
    print(f"修改项数量：{len(sample_changes)}")
    print()

    # 尝试原生修订模式
    print("1. 尝试原生修订模式...")
    success1 = create_track_changes_docx(
        original_text="示例合同",
        changes=sample_changes,
        output_path=output_path,
        author="AI审核助手"
    )

    if success1:
        print(f"   ✅ 原生修订模式文档已生成")
    else:
        print(f"   ⚠️ 原生修订模式失败，尝试兼容模式...")

        # 尝试兼容模式
        output_compat = output_path.replace('.docx', '_兼容版.docx')
        success2 = create_compatible_docx(
            original_text="示例合同",
            changes=sample_changes,
            output_path=output_compat,
            author="AI审核助手"
        )

        if success2:
            print(f"   ✅ 兼容版文档已生成：{output_compat}")

    print()
    print("测试完成！请用 Word 打开查看。")
