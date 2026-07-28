# -*- coding: utf-8 -*-
"""
合同定稿学习模块 (contract_learner.py)
=========================================

功能：
1. 检查定稿文件中的错别字、病句、不通顺表达
2. 学习用户的修改模式和行文风格
3. 保存学习结果供后续审核参考
"""

import re
import json
from datetime import datetime
from typing import List, Dict, Optional
from pathlib import Path
from collections import defaultdict, Counter
from dataclasses import dataclass, field, asdict

from docx import Document
from lxml import etree


@dataclass
class TextIssue:
    issue_type: str
    position: int
    text: str
    suggestion: str
    severity: str
    reason: str


@dataclass
class ModificationPattern:
    pattern_type: str
    original_pattern: str
    user_pattern: str
    frequency: int
    context: str
    examples: List[str]


@dataclass
class WritingStyle:
    preferred_words: List[str]
    sentence_patterns: List[str]
    formal_indicators: List[str]
    clause_structures: List[str]


class TextChecker:
    """文本问题检查器"""

    COMMON_TYPOS = {
        '象': '像', '需': '须', '的': '地', '地': '的',
        '权利': '权力', '定金': '订金', '帐号': '账户',
    }

    LEGAL_TYPOS = {
        '大概': '大约', '如果': '如', '就是': '即',
        '应该': '应当', '必须': '应当', '不能': '无法',
    }

    def __init__(self):
        self.issues = []

    def check_typo(self, text: str) -> List[TextIssue]:
        issues = []
        for old, new in self.COMMON_TYPOS.items():
            if old in text:
                if not self._is_legal_term_context(text, old):
                    pos = text.find(old)
                    issues.append(TextIssue(
                        issue_type='typo',
                        position=pos,
                        text=old,
                        suggestion=new,
                        severity='high',
                        reason=f"可能应为'{new}'（需根据上下文确认）"
                    ))
        return issues

    def _is_legal_term_context(self, text: str, word: str) -> bool:
        legal_terms = ['权利义务', '定金条款', '权利']
        return any(term in text for term in legal_terms if word in term)

    def check_grammar(self, text: str) -> List[TextIssue]:
        issues = []
        lines = text.split('\n')

        for line in lines:
            # 检查过长句子
            if len(line) > 150 and line.count('；') == 0:
                issues.append(TextIssue(
                    issue_type='clarity',
                    position=0,
                    text=line[:50] + '...',
                    suggestion='建议拆分为多个短句',
                    severity='medium',
                    reason='句子过长，建议使用分号拆分'
                ))

            # 检查重复用词
            words = re.findall(r'[\u4e00-\u9fa5]{2,}', line)
            word_counts = Counter(words)
            for word, count in word_counts.items():
                if count >= 3 and len(word) >= 2:
                    issues.append(TextIssue(
                        issue_type='style',
                        position=line.find(word),
                        text=word,
                        suggestion=f'避免重复使用"{word}"',
                        severity='low',
                        reason=f'"{word}"重复使用{count}次'
                    ))

        return issues

    def check_all(self, text: str) -> List[TextIssue]:
        all_issues = []
        all_issues.extend(self.check_typo(text))
        all_issues.extend(self.check_grammar(text))
        return all_issues


class PatternLearner:
    """修改模式学习器"""

    def __init__(self):
        self.patterns = []
        self.writing_style = WritingStyle(
            preferred_words=[],
            sentence_patterns=[],
            formal_indicators=[],
            clause_structures=[]
        )

    def extract_patterns(self, original_text: str, final_text: str) -> List[ModificationPattern]:
        patterns = []

        # 违约金相关修改
        penalty_orig = re.findall(r'违约金[^\n。；]{0,30}', original_text)
        penalty_final = re.findall(r'违约金[^\n。；]{0,30}', final_text)

        for orig in penalty_orig:
            for fin in penalty_final:
                if orig != fin and '违约金' in orig:
                    patterns.append(ModificationPattern(
                        pattern_type='phrase_adjustment',
                        original_pattern=orig,
                        user_pattern=fin,
                        frequency=1,
                        context='违约金条款',
                        examples=[f'{orig} → {fin}']
                    ))

        # 条款数量变化
        orig_clauses = re.findall(r'第[一二三四五六七八九十百零\d]+条', original_text)
        fin_clauses = re.findall(r'第[一二三四五六七八九十百零\d]+条', final_text)

        if len(orig_clauses) != len(fin_clauses):
            patterns.append(ModificationPattern(
                pattern_type='clause_reorganization',
                original_pattern=f'原{len(orig_clauses)}条',
                user_pattern=f'现{len(fin_clauses)}条',
                frequency=1,
                context='条款数量调整',
                examples=[f'条款数: {len(orig_clauses)} → {len(fin_clauses)}']
            ))

        self.patterns = patterns
        return patterns

    def analyze_writing_style(self, text: str):
        formal_words = ['应当', '不得', '依照', '致使', '据此', '故', '即', '均', '届时', '前述']
        found_formal = [w for w in formal_words if w in text]
        self.writing_style.formal_indicators = found_formal

        clause_structures = re.findall(r'[一二三]、|[甲乙丙丁戊]、|（[一二三四]）', text)
        self.writing_style.clause_structures = list(set(clause_structures))

        sentence_patterns = []
        if '；' in text:
            sentence_patterns.append('使用分号连接')
        if '如下' in text or '包括' in text:
            sentence_patterns.append('列举说明')
        self.writing_style.sentence_patterns = sentence_patterns


class ContractLearner:
    """合同定稿学习器"""
    
    # OOXML 命名空间
    NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    def __init__(self, user_id: str = "default"):
        self.user_id = user_id
        self.document_text = ""
        self.document_name = ""
        self.text_checker = TextChecker()
        self.pattern_learner = PatternLearner()
        self.issues = []
        self.patterns = []
        self.writing_style = WritingStyle(
            preferred_words=[],
            sentence_patterns=[],
            formal_indicators=[],
            clause_structures=[]
        )
        # v1.7: XML 修订标记存储
        self.xml_ins_content = []  # 用户新增的内容（<w:ins>）
        self.xml_del_content = []  # 用户删除的内容（<w:del>）

    def extract_track_changes_from_docx(self, docx_path: str) -> Dict:
        """
        从 .docx 文件中提取 OOXML Track Changes 修订标记

        Uses lxml to parse <w:ins> and <w:del> elements.
        This is the core of the "定稿学习" feature described in SKILL.md.

        Args:
            docx_path: .docx 文件路径

        Returns:
            {'insertions': [str, ...], 'deletions': [str, ...]}
        """
        insertions = []
        deletions = []
        
        try:
            doc = Document(docx_path)
            body = doc.element.body
            
            # 提取 <w:ins> 新增内容
            for ins in body.findall('.//w:ins', self.NSMAP):
                texts = []
                for t in ins.findall('.//w:t', self.NSMAP):
                    if t.text:
                        texts.append(t.text)
                full_text = ''.join(texts).strip()
                if full_text:
                    insertions.append(full_text)

            # 提取 <w:del> 删除内容（使用 w:delText）
            for del_elem in body.findall('.//w:del', self.NSMAP):
                texts = []
                for dt in del_elem.findall('.//w:delText', self.NSMAP):
                    if dt.text:
                        texts.append(dt.text)
                full_text = ''.join(texts).strip()
                if full_text:
                    deletions.append(full_text)

            self.xml_ins_content = insertions
            self.xml_del_content = deletions
            
        except Exception as e:
            print(f"提取 Track Changes 失败: {e}")
        
        return {
            'insertions': insertions,
            'deletions': deletions,
        }

    def compare_xml_track_changes(self, original_docx_path: str, revised_docx_path: str) -> Dict:
        """
        XML 层级比对：比较原始审核意见的修订 vs 用户定稿后的修订

        Args:
            original_docx_path: AI 审核意见生成的修订稿
            revised_docx_path: 用户修改后的定稿

        Returns:
            {
                'ai_suggestions': {'insertions': [...], 'deletions': [...]},
                'user_modifications': {'insertions': [...], 'deletions': [...]},
                'adopted': [...],      # 用户采纳的建议
                'modified': [...],     # 用户修改后采纳的
                'rejected': [...],     # 用户未采纳的建议
                'user_additions': [...], # 用户自主新增的修改
            }
        """
        ai_tc = ContractLearner().extract_track_changes_from_docx(original_docx_path)
        user_tc = self.extract_track_changes_from_docx(revised_docx_path)

        ai_ins = ai_tc.get('insertions', [])
        ai_del = ai_tc.get('deletions', [])
        user_ins = user_tc.get('insertions', [])
        user_del = user_tc.get('deletions', [])

        adopted = []
        modified = []
        rejected = []
        user_additions = []

        # 比对 AI 的建议（insertions）是否被用户采纳
        for ai_item in ai_ins:
            matched = False
            for user_item in user_ins:
                # 宽松匹配：用户修改后的文本是否包含 AI 建议的核心内容
                common_words = set(ai_item) & set(user_item)
                overlap = len(common_words) / max(len(set(ai_item)), 1)
                if overlap > 0.5:  # 50% 以上重叠视为采纳
                    if overlap > 0.8:
                        adopted.append({'ai': ai_item[:100], 'user': user_item[:100]})
                    else:
                        modified.append({'ai': ai_item[:100], 'user': user_item[:100]})
                    matched = True
                    break
            if not matched:
                rejected.append(ai_item[:100])

        # 检测用户自主新增的修改（不在 AI 建议中的）
        for user_item in user_ins:
            found_in_ai = False
            for ai_item in ai_ins:
                common = set(user_item) & set(ai_item)
                if len(common) / max(len(set(user_item)), 1) > 0.3:
                    found_in_ai = True
                    break
            if not found_in_ai:
                user_additions.append(user_item[:100])

        return {
            'ai_suggestions': ai_tc,
            'user_modifications': user_tc,
            'adopted': adopted,
            'modified': modified,
            'rejected': rejected,
            'user_additions': user_additions,
        }

    def load_document(self, path: str) -> bool:
        path = Path(path)
        if not path.exists():
            print(f"文件不存在: {path}")
            return False

        self.document_name = path.name

        try:
            if path.suffix == '.docx':
                doc = Document(path)
                self.document_text = '\n'.join([p.text for p in doc.paragraphs])
            elif path.suffix == '.txt':
                with open(path, 'r', encoding='utf-8') as f:
                    self.document_text = f.read()
            return True
        except Exception as e:
            print(f"加载文档失败: {e}")
            return False

    def check_issues(self) -> List[TextIssue]:
        if not self.document_text:
            return []
        self.issues = self.text_checker.check_all(self.document_text)
        return self.issues

    def learn_from_comparison(self, ai_suggestions_path: str = None,
                              final_document_path: str = None) -> List[ModificationPattern]:
        """
        从 AI 建议稿和定稿的对比中学习修改模式
        优先使用 XML 层级比对（Track Changes），回退到文本比对
        """
        # 优先尝试 XML Track Changes 比对
        if ai_suggestions_path and final_document_path:
            try:
                xml_result = self.compare_xml_track_changes(ai_suggestions_path, final_document_path)
                # 将 XML 比对结果转换为 ModificationPattern
                adopted = xml_result.get('adopted', [])
                modified = xml_result.get('modified', [])
                rejected = xml_result.get('rejected', [])
                user_additions = xml_result.get('user_additions', [])

                patterns_from_xml = []
                for item in adopted:
                    patterns_from_xml.append(ModificationPattern(
                        pattern_type='adopted',
                        original_pattern=item.get('ai', ''),
                        user_pattern=item.get('user', ''),
                        frequency=1,
                        context='用户完全采纳的修改',
                        examples=[f'AI: {item.get("ai", "")[:50]} → User: {item.get("user", "")[:50]}']
                    ))
                for item in modified:
                    patterns_from_xml.append(ModificationPattern(
                        pattern_type='modified_adopted',
                        original_pattern=item.get('ai', ''),
                        user_pattern=item.get('user', ''),
                        frequency=1,
                        context='用户调整后采纳的修改',
                        examples=[f'AI: {item.get("ai", "")[:50]} → User: {item.get("user", "")[:50]}']
                    ))
                for item in rejected:
                    patterns_from_xml.append(ModificationPattern(
                        pattern_type='rejected',
                        original_pattern=item,
                        user_pattern='[用户未采纳]',
                        frequency=1,
                        context='用户未采纳的建议',
                        examples=[f'AI suggested: {item[:50]}']
                    ))
                for item in user_additions:
                    patterns_from_xml.append(ModificationPattern(
                        pattern_type='user_original',
                        original_pattern='',
                        user_pattern=item,
                        frequency=1,
                        context='用户自主新增的修改',
                        examples=[f'User added: {item[:50]}']
                    ))

                if patterns_from_xml:
                    self.patterns = patterns_from_xml
                    return self.patterns
            except Exception as e:
                pass  # XML 比对失败，回退到文本比对

        # 回退：文本层面的比对
        ai_text = ""
        if ai_suggestions_path and Path(ai_suggestions_path).exists():
            doc = Document(ai_suggestions_path)
            ai_text = '\n'.join([p.text for p in doc.paragraphs])

        if ai_text:
            self.patterns = self.pattern_learner.extract_patterns(ai_text, self.document_text)
        else:
            self.pattern_learner.analyze_writing_style(self.document_text)
            self.patterns = self.pattern_learner.patterns

        self.writing_style = self.pattern_learner.writing_style
        return self.patterns

    def learn_from_document(self) -> List[ModificationPattern]:
        if not self.document_text:
            return []
        self.pattern_learner.analyze_writing_style(self.document_text)
        self.patterns = self.pattern_learner.patterns
        self.writing_style = self.pattern_learner.writing_style
        return self.patterns

    def format_issues_report(self) -> str:
        if not self.issues:
            return "✅ 未发现明显问题"

        report = ["📋 文本问题报告", "=" * 40]

        by_severity = defaultdict(list)
        for issue in self.issues:
            by_severity[issue.severity].append(issue)

        if by_severity['high']:
            report.append("\n🔴 高优先级问题：")
            for issue in by_severity['high']:
                report.append(f"  • {issue.text} → {issue.suggestion}")

        if by_severity['medium']:
            report.append(f"\n🟡 中优先级问题：{len(by_severity['medium'])} 处")

        if by_severity['low']:
            report.append(f"\n🟢 低优先级：{len(by_severity['low'])} 处")

        return '\n'.join(report)

    def format_patterns_report(self) -> str:
        if not self.patterns:
            return "📝 暂无修改模式数据\n继续使用将自动学习您的修改习惯"

        report = ["📝 定稿学习报告", "=" * 40]

        by_type = defaultdict(list)
        for pattern in self.patterns:
            by_type[pattern.pattern_type].append(pattern)

        # 统计
        adopted_count = len(by_type.get('adopted', []))
        modified_count = len(by_type.get('modified_adopted', []))
        rejected_count = len(by_type.get('rejected', []))
        user_orig_count = len(by_type.get('user_original', []))

        report.append(f"\n📊 学习摘要：")
        report.append(f"  ✅ 完全采纳：{adopted_count} 项")
        report.append(f"  🔄 调整后采纳：{modified_count} 项")
        report.append(f"  ❌ 未采纳：{rejected_count} 项")
        report.append(f"  🆕 自主新增：{user_orig_count} 项")

        # 详细列表
        for ptype, pats in by_type.items():
            if not pats:
                continue
            report.append(f"\n【{ptype}】({len(pats)} 项)")
            shown = set()
            for p in pats:
                key = (p.user_pattern or p.original_pattern)[:30]
                if key and key not in shown:
                    shown.add(key)
                    if p.user_pattern and p.user_pattern != '[用户未采纳]':
                        report.append(f"  → {p.user_pattern[:80]}")
                    if p.original_pattern:
                        report.append(f"  ← {p.original_pattern[:80]}")

        # 如果做了 XML 比对，提示可保存
        if self.xml_ins_content or self.xml_del_content:
            report.append(f"\n💾 XML 修订标记已解析：新增 {len(self.xml_ins_content)} 处，删除 {len(self.xml_del_content)} 处")
            report.append("  可通过 save_learning_data() 持久化本次学习成果")

        return '\n'.join(report)

    def save_learning_data(self, output_path: str):
        data = {
            'user_id': self.user_id,
            'learned_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'document_name': self.document_name,
            'patterns': [asdict(p) for p in self.patterns],
            'writing_style': asdict(self.writing_style),
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        print(f"✅ 学习数据已保存: {output_path}")

    def load_learning_data(self, input_path: str) -> bool:
        if not Path(input_path).exists():
            return False

        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            self.user_id = data.get('user_id', 'default')
            ws_data = data.get('writing_style', {})

            self.writing_style = WritingStyle(
                preferred_words=ws_data.get('preferred_words', []),
                sentence_patterns=ws_data.get('sentence_patterns', []),
                formal_indicators=ws_data.get('formal_indicators', []),
                clause_structures=ws_data.get('clause_structures', [])
            )
            return True
        except Exception as e:
            print(f"加载学习数据失败: {e}")
            return False


def analyze_final_document(
    final_doc_path: str,
    ai_suggestions_path: str = None,
    learning_data_path: str = None
) -> Dict:
    """一键分析定稿文档"""
    learner = ContractLearner()

    if not learner.load_document(final_doc_path):
        return {'error': '无法加载文档'}

    issues = learner.check_issues()

    if ai_suggestions_path:
        patterns = learner.learn_from_comparison(ai_suggestions_path, final_doc_path)
    else:
        patterns = learner.learn_from_document()

    if learning_data_path:
        learner.save_learning_data(learning_data_path)

    return {
        'document_name': learner.document_name,
        'issues_count': len(issues),
        'patterns_count': len(patterns),
        'issues_report': learner.format_issues_report(),
        'patterns_report': learner.format_patterns_report(),
    }


if __name__ == "__main__":
    print("=== 合同定稿学习器测试 ===\n")

    test_text = """
    第一条 产品描述
    甲方向乙方采购服务器设备10台。

    第二条 价款及付款方式
    合同总价为人民币50万元整，甲方应与收到货物后30日内付款。

    第三条 违约责任
    任何一方违约因当承担相应法律责任，违约金为合同总价的50%。
    """

    checker = TextChecker()
    issues = checker.check_all(test_text)
    print(f"发现 {len(issues)} 处问题:")
    for issue in issues:
        print(f"  - [{issue.issue_type}] {issue.text} → {issue.suggestion}")
