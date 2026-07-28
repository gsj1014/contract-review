# -*- coding: utf-8 -*-
"""
合同修订稿生成模块 (contract_revision.py)
==========================================

根据合同审核意见，自动生成修订稿。

功能：
1. 基于审核意见生成修改建议
2. 标记添加/删除内容
3. 生成 Word 修订模式文档（.docx）

使用示例：
    from contract_revision import ContractReviser

    revisor = ContractReviser(contract_text, review_result)
    revised = revisor.generate_marked_text()  # 获取带标记的文本

    # 生成 Word 修订文档
    revisor.save_as_docx("修订稿.docx")
"""

import re
import sys
import os
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from lxml import etree
import datetime

# 从 export_report 导入示范条款库
try:
    from .export_report import MODEL_CLAUSES, generate_model_clause, _match_model_clause_key
except ImportError:
    sys.path.insert(0, os.path.dirname(__file__))
    from export_report import MODEL_CLAUSES, generate_model_clause, _match_model_clause_key


def clip_text(text: str, max_len: int = 200) -> str:
    """截断文本，保持可读性"""
    if len(text) <= max_len:
        return text
    return text[:max_len] + '...'


class ChangeType(Enum):
    """修改类型枚举"""
    ADD = "add"           # 新增内容
    DELETE = "delete"     # 删除内容
    REPLACE = "replace"    # 替换内容
    COMMENT = "comment"    # 仅添加批注


@dataclass
class ContractChange:
    """合同修改项"""
    # 修改位置
    clause_title: str      # 所属条款（如"第三条 价款"）
    original_text: str     # 原文
    new_text: str          # 修改后文本

    # 修改类型
    change_type: ChangeType

    # 修改依据
    reason: str            # 修改原因（审核意见）
    legal_basis: str       # 法律依据

    # 位置信息
    line_number: int = 0   # 行号（原始文本中的位置）

    def __str__(self):
        if self.change_type == ChangeType.ADD:
            return f"【新增】{self.clause_title}: {self.new_text}"
        elif self.change_type == ChangeType.DELETE:
            return f"【删除】{self.clause_title}: {self.original_text}"
        elif self.change_type == ChangeType.REPLACE:
            return f"【修改】{self.clause_title}: {self.original_text} → {self.new_text}"
        else:
            return f"【批注】{self.clause_title}: {self.reason}"


@dataclass
class RevisionResult:
    """修订结果"""
    original_text: str              # 原始文本
    revised_text: str               # 修订后文本（带标记）
    changes: List[ContractChange]   # 修改项列表
    statistics: Dict[str, int]     # 统计信息


class ContractReviser:
    """
    合同修订器

    根据审核意见，自动生成合同修订稿。

    使用方式：
        revisor = ContractReviser(
            contract_text=original,
            review_result=review_result,  # 审核报告结果
            risk_items=risk_items,         # 风险条款列表
            gap_items=gap_items,           # 差异分析列表
        )

        # 获取带标记的修订文本
        marked_text = revisor.generate_marked_text()

        # 保存为 Word 文档
        revisor.save_as_docx("修订稿.docx")
    """

    def __init__(
        self,
        contract_text: str,
        review_result: Dict = None,
        risk_items: List[Dict] = None,
        gap_items: List[Dict] = None,
        legal_items: List[Dict] = None,
    ):
        """
        初始化合同修订器

        Args:
            contract_text: 原始合同文本
            review_result: 审核报告结果（可选）
            risk_items: 风险条款列表（可选）
            gap_items: 差异分析列表（可选）
            legal_items: 法律合规问题列表（可选）
        """
        self.original_text = contract_text
        self.review_result = review_result or {}
        self.risk_items = risk_items or []
        self.gap_items = gap_items or []
        self.legal_items = legal_items or []

        # 解析合同结构
        self.clauses = self._parse_clauses()

        # 生成修改项
        self.changes = self._generate_changes()

    def _parse_clauses(self) -> List[Dict]:
        """解析合同条款"""
        clauses = []
        lines = self.original_text.split('\n')

        current_clause = None
        current_content = []
        current_line = 0

        for i, line in enumerate(lines):
            # 检测条款标题
            clause_patterns = [
                r'^第[一二三四五六七八九十百零\d]+条\s*[：:]\s*(.+)',
                r'^[第一二三四五六七八九十]+[、.]\s*(.+)',
                r'^第[一二三四五六七八九十百零\d]+条\s*(.+)',
            ]

            matched = False
            for pattern in clause_patterns:
                match = re.match(pattern, line.strip())
                if match:
                    # 保存前一个条款
                    if current_clause:
                        clauses.append({
                            'title': current_clause,
                            'content': '\n'.join(current_content),
                            'start_line': current_line,
                        })

                    current_clause = line.strip()
                    current_content = []
                    current_line = i
                    matched = True
                    break

            if not matched and current_clause:
                current_content.append(line)

        # 保存最后一个条款
        if current_clause:
            clauses.append({
                'title': current_clause,
                'content': '\n'.join(current_content),
                'start_line': current_line,
            })

        return clauses

    def _find_clause_for_risk(self, item: Dict) -> Optional[Dict]:
        """根据风险项找到对应的条款"""
        item_text = str(item)

        for clause in self.clauses:
            clause_text = clause['title'] + clause['content']

            # 简单的关键词匹配
            keywords = self._extract_keywords(item.get('title', item.get('description', '')))
            for kw in keywords:
                if kw in clause_text:
                    return clause

        return None

    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        keywords = []
        text = text.replace('条款', '').replace('合同', '')

        # 提取常见关键概念
        concepts = [
            '标的', '质量', '数量', '价款', '付款', '支付', '交付',
            '验收', '违约', '责任', '风险', '所有', '权利', '义务',
            '争议', '解决', '解除', '终止', '变更', '转让', '保密',
            '不可抗力', '通知', '适用', '生效', '无效', '效力'
        ]

        for concept in concepts:
            if concept in text:
                keywords.append(concept)

        return keywords

    def _generate_changes(self) -> List[ContractChange]:
        """根据审核意见生成修改项"""
        changes = []

        # 处理风险条款
        for item in self.risk_items:
            change = self._generate_change_from_risk(item)
            if change:
                changes.append(change)

        # 处理差异分析
        for item in self.gap_items:
            change = self._generate_change_from_gap(item)
            if change:
                changes.append(change)

        # 处理法律合规问题
        for item in self.legal_items:
            change = self._generate_change_from_legal(item)
            if change:
                changes.append(change)

        return changes

    def _generate_change_from_risk(self, item: Dict) -> Optional[ContractChange]:
        """从风险条款生成修改（v2.0：接入 MODEL_CLAUSES，覆盖全部风险类型）"""
        clause = self._find_clause_for_risk(item)
        title = item.get('title', '')
        description = item.get('description', '')
        suggestion = item.get('suggestion', '')

        # 1. 优先从 MODEL_CLAUSES 获取示范条款
        model = generate_model_clause(title)
        concise = model.get('concise', '')
        full = model.get('full', '')

        # 2. 判断修改类型
        is_missing = any(kw in title for kw in ['缺少', '缺失', '无验收', '无违约解约权'])
        is_high_risk = any(kw in title for kw in ['可能导致', '过高', '不利', '不明', '过宽', '过低', '缺失'])

        if not clause:
            # 找不到对应条款 = 缺失项，用 ADD
            if is_missing or not clause:
                return ContractChange(
                    clause_title=title,
                    original_text='【未找到相关条款】',
                    new_text=full or concise or suggestion,
                    change_type=ChangeType.ADD,
                    reason=description or title,
                    legal_basis=item.get('reference', ''),
                )

        # 3. 有对应条款 → 基于 MODEL_CLAUSES 生成 REPLACE/ADD
        if is_missing:
            return ContractChange(
                clause_title=clause['title'],
                original_text=clause['content'][:200] if clause['content'].strip() else '【条款内容缺失】',
                new_text=full or concise or suggestion,
                change_type=ChangeType.ADD,
                reason=description or title,
                legal_basis=item.get('reference', ''),
            )

        # 4. 普通风险条款 → REPLACE（用全面版示范条款替换原文）
        addition_text = full or concise or suggestion
        original_text = clip_text(self._extract_original_text(clause, title), 200)
        
        return ContractChange(
            clause_title=clause['title'],
            original_text=original_text,
            new_text=addition_text,
            change_type=ChangeType.REPLACE,
            reason=description or title,
            legal_basis=item.get('reference', ''),
        )

    def _generate_change_from_gap(self, item: Dict) -> Optional[ContractChange]:
        """从差异分析生成修改（v2.0：接入 MODEL_CLAUSES）"""
        gap_type = item.get('type', '')
        clause_name = item.get('clause', '')
        suggestion = item.get('suggestion', '')
        legal_basis = item.get('legal_basis', '')

        # 从 MODEL_CLAUSES 获取示范条款
        model = generate_model_clause(clause_name)
        addition_text = model.get('full') or model.get('concise') or suggestion

        # 查找最近位置
        for clause in self.clauses:
            if clause_name in clause['title'] or clause_name in clause['content']:
                return ContractChange(
                    clause_title=clause['title'],
                    original_text='【建议在此处增加条款】',
                    new_text=addition_text,
                    change_type=ChangeType.ADD if gap_type == '缺失' else ChangeType.REPLACE,
                    reason=f"{gap_type}条款：{clause_name}",
                    legal_basis=legal_basis
                )

        # 找不到匹配条款 → 作为独立新增条款放到末尾
        return ContractChange(
            clause_title=f'【新增】{clause_name}',
            original_text='【无原条款】',
            new_text=addition_text,
            change_type=ChangeType.ADD,
            reason=f"{gap_type}条款：{clause_name}",
            legal_basis=legal_basis
        )

    def _generate_change_from_legal(self, item: Dict) -> Optional[ContractChange]:
        """从法律合规问题生成修改（v2.0：接入 MODEL_CLAUSES）"""
        clause_name = item.get('clause', '')
        suggestion = item.get('suggestion', '')
        legal_basis = item.get('legal_basis', '')

        model = generate_model_clause(clause_name)
        addition_text = model.get('full') or model.get('concise') or suggestion

        for clause in self.clauses:
            if clause_name in clause['title'] or clause_name in clause['content']:
                return ContractChange(
                    clause_title=clause['title'],
                    original_text=clause['content'][:300],
                    new_text=addition_text,
                    change_type=ChangeType.REPLACE,
                    reason=f"法律合规问题：{clause_name}",
                    legal_basis=legal_basis
                )

        return ContractChange(
            clause_title=clause_name,
            original_text='【需依法补充】',
            new_text=addition_text,
            change_type=ChangeType.ADD,
            reason=f"法律合规：{clause_name}",
            legal_basis=legal_basis
        )

    def _extract_original_text(self, clause: Dict, risk_title: str) -> str:
        """根据风险类型从条款内容中提取相关的原文片段"""
        content = clause['content']
        keyword_map = {
            '管辖': '管辖',
            '争议': '争议',
            '付款': '付款',
            '支付': '支付',
            '违约金': '违约金',
            '违约': '违约',
            '不可抗力': '不可抗力',
            '保密': '保密',
            '知识产权': '知识产权',
            '验收': '验收',
            '解除': '解除',
        }
        for key, kw in keyword_map.items():
            if key in risk_title and kw in content:
                idx = content.find(kw)
                start = max(0, idx - 20)
                end = min(len(content), idx + 100)
                return content[start:end].strip()
        return content[:200].strip()

    def generate_marked_text(self) -> str:
        """生成带标记的修订文本"""
        if not self.changes:
            return self.original_text

        result = self.original_text
        revision_markers = []

        for i, change in enumerate(self.changes, 1):
            marker = self._format_change_marker(change, i)
            revision_markers.append(marker)

        # 添加修订标记
        header = f"""
{'='*60}
合同修订稿（Track Changes）
{'='*60}

修订说明：
1. 【删除线】表示原文需要删除的内容
2. 【下划线】表示新增的内容
3. 修改项编号对应下方详细说明

{'='*60}

"""

        changes_detail = "\n".join(revision_markers)

        footer = f"""

{'='*60}
修订统计
{'='*60}

- 新增条款：{sum(1 for c in self.changes if c.change_type == ChangeType.ADD)} 项
- 修改条款：{sum(1 for c in self.changes if c.change_type == ChangeType.REPLACE)} 项
- 删除条款：{sum(1 for c in self.changes if c.change_type == ChangeType.DELETE)} 项
- 合计修改：{len(self.changes)} 项

{'='*60}
"""

        return header + result + "\n\n" + changes_detail + footer

    def _format_change_marker(self, change: ContractChange, index: int) -> str:
        """格式化修改标记"""
        marker = f"""
--- 修改项 {index} ---
所属条款：{change.clause_title}

"""

        if change.change_type == ChangeType.ADD:
            marker += f"【新增】\n{change.new_text}\n"
        elif change.change_type == ChangeType.DELETE:
            marker += f"【删除】\n{change.original_text}\n"
        elif change.change_type == ChangeType.REPLACE:
            marker += f"【修改前】\n{change.original_text}\n\n【修改后】\n{change.new_text}\n"

        marker += f"""
修改原因：{change.reason}
法律依据：{change.legal_basis if change.legal_basis else '无' }"""

        return marker

    def generate_statistics(self) -> Dict[str, int]:
        """生成修订统计"""
        return {
            'total_changes': len(self.changes),
            'additions': sum(1 for c in self.changes if c.change_type == ChangeType.ADD),
            'deletions': sum(1 for c in self.changes if c.change_type == ChangeType.DELETE),
            'replacements': sum(1 for c in self.changes if c.change_type == ChangeType.REPLACE),
            'comments': sum(1 for c in self.changes if c.change_type == ChangeType.COMMENT),
        }

    def save_as_docx(self, output_path: str):
        """
        保存为 Word 修订文档（原生 OOXML Track Changes）

        使用 <w:ins> 标记新增、<w:del> 标记删除，
        Word 打开后可正确显示修订模式并支持"接受/拒绝修订"。

        Args:
            output_path: 输出文件路径
        """
        doc = Document()

        # 设置默认段落样式
        style = doc.styles['Normal']
        style.font.size = Pt(11)
        style.font.name = '宋体'

        # 标题
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('合同修订稿（Track Changes）')
        title_run.bold = True
        title_run.font.size = Pt(16)

        doc.add_paragraph()

        # 修订统计
        stats = self.generate_statistics()
        stats_para = doc.add_paragraph()
        stats_para.add_run(f'修订统计：新增 {stats["additions"]} 项、修改 {stats["replacements"]} 项、删除 {stats["deletions"]} 项，合计 {stats["total_changes"]} 项').font.size = Pt(9)
        doc.add_paragraph()

        # 导出原始合同结构，并在需要的位置插入修订
        # 按条款组织修订
        changes_by_clause = {}
        for change in self.changes:
            title = change.clause_title
            if title not in changes_by_clause:
                changes_by_clause[title] = []
            changes_by_clause[title].append(change)

        # 遍历子条款，按顺序输出
        for clause in self.clauses:
            clause_title = clause['title']
            content = clause['content']

            # 条款标题
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(clause_title)
            heading_run.bold = True
            heading_run.font.size = Pt(11)

            # 该条款下的修改项
            clause_changes = []
            for title_key, changes in changes_by_clause.items():
                if title_key == clause_title or title_key in clause_title or clause_title in title_key:
                    clause_changes.extend(changes)

            if clause_changes:
                for change in clause_changes:
                    if change.change_type == ChangeType.DELETE:
                        # 删除：显示原文为红色删除线（delText）
                        self._add_track_change_del(doc, change.original_text, change.reason)
                    elif change.change_type == ChangeType.REPLACE:
                        # 替换：先删除原文，再插入新文
                        self._add_track_change_del(doc, change.original_text, 
                                                   f"【删除原因】{change.reason}")
                        self._add_track_change_ins(doc, change.new_text,
                                                   f"【修改原因】{change.reason}")
                    elif change.change_type == ChangeType.ADD:
                        # 新增：蓝色下划线
                        self._add_track_change_ins(doc, change.new_text,
                                                   f"【新增原因】{change.reason}")
            else:
                # 无修改的条款，按原样输出
                for line in content.split('\n'):
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)

            doc.add_paragraph()  # 条款间空行

        # 详细修改说明（独立章节）
        doc.add_page_break()
        doc.add_heading('详细修改说明', level=1)

        for i, change in enumerate(self.changes, 1):
            self._add_change_detail_to_doc(doc, change, i)

        # 保存
        doc.save(output_path)

    def _add_track_change_ins(self, doc: Document, text: str, comment: str):
        """添加一个 <w:ins> 修订标记段落（原生 OOXML Track Changes 插入）

        Args:
            doc: Document 对象
            text: 要插入的文本
            comment: 修订批注（作为段落后的说明）
        """
        para = doc.add_paragraph()
        p_elem = para._element

        # 创建 <w:ins> 元素
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), str(len(doc.paragraphs)))  # 唯一ID
        ins.set(qn('w:author'), '合同侠·WorkBuddy')
        ins.set(qn('w:date'), datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))

        # 在 <w:ins> 内创建 <w:r> (run)
        r = OxmlElement('w:r')

        # 设置蓝色 + 下划线样式（Word 修订模式的新增标准样式）
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')  # 蓝色
        rPr.append(color)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')  # 单下划线
        rPr.append(u)
        r.append(rPr)

        # 添加文本
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)

        ins.append(r)
        p_elem.append(ins)

    def _add_track_change_del(self, doc: Document, text: str, comment: str):
        """添加一个 <w:del> 修订标记段落（原生 OOXML Track Changes 删除）

        Args:
            doc: Document 对象
            text: 要标记为删除的文本
            comment: 删除原因
        """
        para = doc.add_paragraph()
        p_elem = para._element

        # 创建 <w:del> 元素
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), str(len(doc.paragraphs) + 1000))  # 唯一ID（偏移避免冲突）
        del_elem.set(qn('w:author'), '合同侠·WorkBuddy')
        del_elem.set(qn('w:date'), datetime.datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))

        # 在 <w:del> 内创建 <w:r>
        r = OxmlElement('w:r')

        # 设置红色 + 删除线样式
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'FF0000')  # 红色
        rPr.append(color)
        strike = OxmlElement('w:strike')
        strike.set(qn('w:val'), 'true')
        rPr.append(strike)
        r.append(rPr)

        # 添加删除文本（使用 <w:delText> 而非 <w:t>）
        dt = OxmlElement('w:delText')
        dt.set(qn('xml:space'), 'preserve')
        dt.text = text
        r.append(dt)

        del_elem.append(r)
        p_elem.append(del_elem)

    def _add_change_to_doc(self, doc: Document, change: ContractChange):
        """添加修改内容到 Word 文档（兼容旧接口）"""
        self._add_track_change_ins(doc, change.new_text, change.reason)

    def _add_change_detail_to_doc(self, doc: Document, change: ContractChange, index: int):
        """添加详细修改说明到 Word 文档"""
        # 修改项标题
        heading = doc.add_heading(f'修改项 {index}', level=2)

        # 所属条款
        p1 = doc.add_paragraph()
        p1.add_run('所属条款：').bold = True
        p1.add_run(change.clause_title)

        # 修改类型
        p2 = doc.add_paragraph()
        p2.add_run('修改类型：').bold = True
        p2.add_run({
            ChangeType.ADD: '新增',
            ChangeType.DELETE: '删除',
            ChangeType.REPLACE: '替换',
            ChangeType.COMMENT: '批注'
        }.get(change.change_type, '未知'))

        # 修改原因
        p3 = doc.add_paragraph()
        p3.add_run('修改原因：').bold = True
        p3.add_run(change.reason)

        # 法律依据
        if change.legal_basis:
            p4 = doc.add_paragraph()
            p4.add_run('法律依据：').bold = True
            p4.add_run(change.legal_basis)

        doc.add_paragraph()


# 便捷函数
def generate_revision(
    contract_text: str,
    review_result: Dict = None,
    risk_items: List[Dict] = None,
    gap_items: List[Dict] = None,
    legal_items: List[Dict] = None,
    output_path: str = None
) -> str:
    """
    生成合同修订稿

    参数:
        contract_text: 原始合同文本
        review_result: 审核报告结果
        risk_items: 风险条款列表
        gap_items: 差异分析列表
        legal_items: 法律合规问题列表
        output_path: 输出文件路径（可选）

    返回:
        带标记的修订文本
    """
    revisor = ContractReviser(
        contract_text=contract_text,
        review_result=review_result,
        risk_items=risk_items,
        gap_items=gap_items,
        legal_items=legal_items,
    )

    marked_text = revisor.generate_marked_text()

    if output_path:
        revisor.save_as_docx(output_path)

    return marked_text


# 测试代码
if __name__ == "__main__":
    sample_contract = """
采购合同

甲方（买方）：大连某科技有限公司
乙方（卖方）：深圳某供应商

第一条 产品描述
甲方向乙方采购服务器设备10台。

第二条 价款
合同总价为人民币50万元整。

第三条 付款条件
甲方应于收货后30日内付款。

第四条 交付时间
乙方应于合同签订后15日内交付产品。

第五条 违约责任
任何一方违约应承担相应法律责任，违约金为合同总价的50%。

第六条 争议解决
因本合同引起的争议，提交甲方所在地人民法院管辖。
"""

    # 模拟审核意见
    risk_items = [
        {
            'title': '违约金比例过高',
            'description': '合同约定违约金为合同总价的50%，可能被认定为过高',
            'suggestion': '建议将违约金比例调整至不超过合同金额的30%为宜',
            'risk_level': 'high'
        },
        {
            'title': '付款条款不完善',
            'description': '付款条款缺少付款方式和发票相关约定',
            'suggestion': '建议补充付款方式和发票条款',
            'risk_level': 'medium'
        },
    ]

    print("=== 合同修订稿生成测试 ===\n")

    revisor = ContractReviser(
        contract_text=sample_contract,
        risk_items=risk_items
    )

    print(f"检测到 {len(revisor.changes)} 项需要修改\n")

    marked_text = revisor.generate_marked_text()
    print(marked_text)

    # 保存为 Word
    try:
        import tempfile, os
        output_path = os.path.join(tempfile.gettempdir(), '修订稿测试.docx')
        revisor.save_as_docx(output_path)
        print(f"\n✅ Word 文档已保存到: {output_path}")
    except Exception as e:
        print(f"\n保存 Word 失败: {e}")
