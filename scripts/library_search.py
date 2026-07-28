#!/usr/bin/env python3
"""
合同库检索模块 - Phase 2
功能：在合同库中检索与待审核合同最相似的模板
"""

import os
import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass

# 合同库根目录
TEMPLATE_BASE = Path(__file__).parent.parent / "references" / "template-library"


@dataclass
class TemplateMatch:
    """模板匹配结果"""
    template_name: str       # 模板文件名
    template_path: str       # 模板路径
    similarity_score: float  # 相似度得分 (0-100)
    matched_keywords: List[str]  # 匹配的关键词
    preview: str            # 模板内容预览（前200字）


class ContractLibrarySearch:
    """合同库检索引擎"""

    # 核心条款类型及其关键词
    CLAUSE_TYPES = {
        '基本信息': ['甲方', '乙方', '丙方', '丁方', '合同编号', '签署日期', '签订地点'],
        '定义条款': ['定义', '术语', '“', '”'],
        '标的条款': ['标的', '产品', '服务', '货物', '商品', '工作成果'],
        '价款条款': ['价款', '价格', '金额', '费用', '报酬', '付款', '支付', '结算', '税费'],
        '交付条款': ['交付', '交付时间', '交付地点', '交付方式', '验收', '检验', '签收'],
        '质量条款': ['质量', '规格', '标准', '瑕疵', '缺陷', '保修', '质保'],
        '权利义务': ['权利', '义务', '责任', '保证', '承诺'],
        '违约条款': ['违约', '违约金', '赔偿', '损失', '救济'],
        '保密条款': ['保密', '机密', '信息披露', '商业秘密'],
        '知识产权条款': ['知识产权', '专利', '商标', '著作权', '版权', '权属', '授权', '许可'],
        '不可抗力': ['不可抗力', '自然灾害', '政府行为', '情势变更'],
        '争议解决': ['争议', '仲裁', '诉讼', '管辖', '法院', '法律适用'],
        '终止条款': ['终止', '解除', '期满', '届满', '失效', '续期'],
        '通知条款': ['通知', '送达', '地址', '联系方式', '变更'],
        '其他条款': ['附件', '份数', '生效', '修改', '补充', '其他'],
    }

    def __init__(self, template_base: str = None):
        """
        初始化检索引擎

        Args:
            template_base: 模板库根目录路径
        """
        self.template_base = Path(template_base) if template_base else TEMPLATE_BASE

    def search(self, contract_text: str, contract_type: str = None,
               top_k: int = 5) -> List[TemplateMatch]:
        """
        检索最相似的模板

        Args:
            contract_text: 待审核合同文本
            contract_type: 可选的合同类型（缩小搜索范围）
            top_k: 返回前k个最相似的结果

        Returns:
            模板匹配结果列表
        """
        # 1. 确定搜索范围
        search_dirs = self._get_search_dirs(contract_type)

        # 2. 加载所有候选模板
        candidates = self._load_candidates(search_dirs)

        # 3. 计算相似度
        scored = []
        for candidate in candidates:
            score, matched_kw = self._calculate_similarity(contract_text, candidate)
            if score > 0:
                scored.append({
                    'template': candidate,
                    'score': score,
                    'matched_keywords': matched_kw,
                })

        # 4. 排序并返回top_k
        scored.sort(key=lambda x: x['score'], reverse=True)
        results = []
        for item in scored[:top_k]:
            results.append(TemplateMatch(
                template_name=item['template']['name'],
                template_path=item['template']['path'],
                similarity_score=item['score'],
                matched_keywords=item['matched_keywords'],
                preview=item['template']['content'][:300] + '...' if len(item['template']['content']) > 300 else item['template']['content']
            ))

        return results

    def _get_search_dirs(self, contract_type: str = None) -> List[Path]:
        """获取待搜索的目录"""
        if contract_type and (self.template_base / contract_type).exists():
            return [self.template_base / contract_type]

        # 搜索所有目录
        dirs = []
        if self.template_base.exists():
            for item in self.template_base.iterdir():
                if item.is_dir() and not item.name.startswith('.'):
                    dirs.append(item)
        return dirs

    def _load_candidates(self, dirs: List[Path]) -> List[Dict]:
        """加载候选模板 - 支持 .md 和 .docx 文件"""
        candidates = []
        for dir_path in dirs:
            contract_type = dir_path.name
            
            # 加载 .md 文件
            for md_file in dir_path.glob("*.md"):
                if md_file.name.startswith('_'):
                    continue  # 跳过占位模板
                try:
                    with open(md_file, 'r', encoding='utf-8') as f:
                        content = f.read()
                    candidates.append({
                        'name': md_file.stem,
                        'path': str(md_file),
                        'type': contract_type,
                        'content': content,
                        'format': 'md',
                    })
                except Exception:
                    continue
            
            # 加载 .docx 文件
            for docx_file in dir_path.glob("*.docx"):
                if docx_file.name.startswith('_'):
                    continue
                try:
                    content = self._read_docx_content(docx_file)
                    if content:
                        candidates.append({
                            'name': docx_file.stem,
                            'path': str(docx_file),
                            'type': contract_type,
                            'content': content,
                            'format': 'docx',
                        })
                except Exception:
                    continue
        
        return candidates
    
    def _read_docx_content(self, docx_path: Path) -> str:
        """读取 .docx 文件内容"""
        try:
            from docx import Document
            doc = Document(docx_path)
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return '\n'.join(paragraphs)
        except ImportError:
            print("python-docx 未安装，无法读取 .docx 文件")
            return ""
        except Exception as e:
            print(f"读取 .docx 文件失败: {e}")
            return ""

    def _calculate_similarity(self, text: str, template: Dict) -> Tuple[float, List[str]]:
        """
        计算相似度 - 增强版算法

        Returns:
            (相似度得分, 匹配的关键词列表)
        """
        text_lower = text.lower()
        content = template['content'].lower()

        matched_keywords = []
        total_score = 0.0

        # 1. 条款类型匹配（高权重）
        for clause_type, keywords in self.CLAUSE_TYPES.items():
            type_score = 0
            clause_found_in_text = 0
            clause_found_in_template = 0
            
            for kw in keywords:
                if kw in text_lower:
                    clause_found_in_text += 1
                if kw in content:
                    clause_found_in_template += 1
                if kw in text_lower and kw in content:
                    type_score += 1
                    if kw not in matched_keywords:
                        matched_keywords.append(kw)
            
            # 条款类型匹配得分 = 该类型关键词在两者中都出现的比例
            if keywords:
                match_ratio = type_score / len(keywords)
                # 如果模板中有这个条款类型但合同中没有，扣分
                if clause_found_in_template > 0 and clause_found_in_text == 0:
                    match_ratio = -0.2  # 模板有但合同没有，惩罚
                total_score += match_ratio * 15  # 条款匹配权重提高

        # 2. N-gram 相似度（句子级别）
        text_ngrams = self._extract_ngrams(text_lower, 3)
        template_ngrams = self._extract_ngrams(content, 3)
        
        common_ngrams = set(text_ngrams) & set(template_ngrams)
        if text_ngrams and template_ngrams:
            ngram_similarity = len(common_ngrams) / max(len(text_ngrams), len(template_ngrams))
            total_score += ngram_similarity * 30

        # 3. 关键词密度相似度
        text_len = len(text)
        template_len = len(template['content'])

        if template_len > 0:
            density = len(matched_keywords) / max(template_len / 1000, 1)  # 每千字匹配数
            total_score += density * 10

        # 4. 整体相似度（TF-IDF 简化版）
        text_words = set(re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9]{2,}', text_lower))
        template_words = set(re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9]{2,}', content))
        
        common_words = text_words & template_words
        if text_words:
            word_similarity = len(common_words) / len(text_words) * 15
            total_score += word_similarity

        # 5. 合同类型匹配加成
        if template['type'] in text_lower:
            total_score += 5

        return min(total_score, 100), matched_keywords[:15]  # 最高100分

    def _extract_ngrams(self, text: str, n: int = 3) -> set:
        """提取 N-gram 特征"""
        # 中文字符处理
        chars = [c for c in text if c.strip() or c in '，。；！？']
        ngrams = set()
        for i in range(len(chars) - n + 1):
            ngram = ''.join(chars[i:i+n])
            # 过滤纯标点符号
            if not all(c in '，。；！？、（）【】""''：' for c in ngram):
                ngrams.add(ngram)
        return ngrams

    def get_template_content(self, template_path: str) -> str:
        """获取模板完整内容"""
        path = Path(template_path)
        
        if not path.exists():
            return ""
        
        try:
            # 处理 .md 文件
            if path.suffix == '.md':
                with open(path, 'r', encoding='utf-8') as f:
                    return f.read()
            # 处理 .docx 文件
            elif path.suffix == '.docx':
                return self._read_docx_content(path)
            else:
                return ""
        except Exception:
            return ""

    def list_templates_by_type(self, contract_type: str) -> List[str]:
        """列出指定类型的全部模板"""
        type_dir = self.template_base / contract_type
        if not type_dir.exists():
            return []
        return [f.stem for f in type_dir.glob("*.md") if not f.stem.startswith('_')]


def main(contract_text: str, contract_type: str = None) -> List[Dict]:
    """
    主入口函数

    Args:
        contract_text: 待审核合同文本
        contract_type: 可选的合同类型

    Returns:
        匹配结果列表
    """
    search = ContractLibrarySearch()
    results = search.search(contract_text, contract_type)

    return [
        {
            'name': r.template_name,
            'path': r.template_path,
            'score': r.similarity_score,
            'matched_keywords': r.matched_keywords,
            'preview': r.preview,
        }
        for r in results
    ]


if __name__ == '__main__':
    # 测试
    sample = """
    咨询服务合同

    甲方（买方）：某科技有限公司
    乙方（卖方）：某供应商

    第一条 产品描述
    甲方向乙方采购以下产品：服务器设备10台。

    第二条 价格及付款
    合同总价为人民币100万元，甲方应于收货后30日内付款。

    第三条 交付及验收
    乙方应于合同签订后15日内交付产品，甲方验收合格后签收。

    第四条 违约责任
    任何一方违约应承担相应法律责任。
    """

    results = main(sample, '买卖合同')
    import json
    print(json.dumps(results, ensure_ascii=False, indent=2))
