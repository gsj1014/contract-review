# -*- coding: utf-8 -*-
"""
ledger_to_revision.py — 台账驱动的确定性修订稿生成（hybrid 架构）
================================================================

把 review-issues-contract.json 账本中 status=confirmed 的问题，
收成一份结构化 render plan，再一次性确定性地应用到原 docx，
带修订收敛 + 归档。

设计取舍（相比 contract-copilot 全量 plan-first）：
- 不做全量 plan-first：AI 仍在审核阶段直接看真实 XML 做四段式判断，保持精度
- 仅在「生成修订稿」边界引入 plan：账本 confirmed → render plan → 确定性应用
- 账本已 80% 是 plan，本脚本只补：action 类型 / 精确 replace 目标文本 / needs_negotiation
- 原 docx 上做「精确 in-place 文本替换」，而非从纯文本重建新文档（区别于 contract_revision.py）

落痕分流规则（来自 SKILL.md「修订落痕分流规则」）：
- 直接修订 revise：confirmed 且非需谈判 → 原文件加 <w:ins>/<w:del>
- 仅批注 comment：confirmed 但 needs_negotiation=true（高风险/商业取舍）→ 不落痕，仅进意见书
- 仅进意见书 opinion：用户选正式意见书 / 提示性低风险 → 不触碰原文件

edit_policy 三档：revise-first / balanced(默认) / comment-first
  - revise-first：能改尽量改，仅真正需谈判转批注
  - balanced：低/中风险直接改，高风险需谈判转批注+进意见书
  - comment-first：全部先批注不落痕，用户逐条确认后再生成

用法：
  python3 ledger_to_revision.py \
      --ledger .workbuddy/review-issues-contract.json \
      --doc 原始合同.docx \
      --out 原始合同_修订版_2026-07-13.docx \
      --policy balanced \
      --plan-out review-plan.json \
      --author "合同智囊·WorkBuddy"

依赖：python-docx + lxml（与本 Skill 其余脚本一致）
"""

import argparse
import copy
import datetime
import json
import os
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "http://www.w3.org/XML/1998/namespace"


# ============================================================
# 账本加载与 render plan 构建
# ============================================================

def load_ledger(path):
    """读取审核问题台账 JSON"""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def build_render_plan(ledger, policy="balanced"):
    """
    把账本中 status=confirmed 的问题收成 render plan。

    每条 plan item 字段：
      id, chapter, section, risk_level, suggestion_type, needs_negotiation,
      action (replace/insert/delete/comment), old, new, anchor
    """
    plan = []
    issues = ledger.get("issues", [])

    for iss in issues:
        if iss.get("status") != "confirmed":
            continue  # 落痕门禁：只认 confirmed

        risk = iss.get("risk_level", "低")
        sug = iss.get("suggestion_type", "modify")
        needs_neg = iss.get("needs_negotiation", risk == "高")
        simple = (iss.get("model_clause_simple") or "").strip()
        full = (iss.get("model_clause_full") or "").strip()
        old = (iss.get("original_text") or "").strip()

        item = {
            "id": iss.get("id"),
            "chapter": iss.get("chapter"),
            "section": iss.get("section"),
            "risk_level": risk,
            "suggestion_type": sug,
            "needs_negotiation": needs_neg,
        }

        # comment-first 策略：全部先批注不落痕
        if policy == "comment-first":
            item["action"] = "comment"
            item["new"] = full or simple
            plan.append(item)
            continue

        # 需谈判（高风险 / 商业取舍 / 改变交易结构）-> 仅批注，不落痕
        # 注：needs_negotiation 由审核阶段写入账本；缺省时按 risk_level=="高" 推断
        if needs_neg:
            item["action"] = "comment"
            item["new"] = full or simple
            plan.append(item)
            continue

        if sug == "add":
            item["action"] = "insert"
            item["new"] = full or simple
            item["anchor"] = iss.get("section") or iss.get("chapter") or ""
            plan.append(item)
            continue

        if sug == "delete":
            item["action"] = "delete"
            item["old"] = old
            plan.append(item)
            continue

        # modify（默认）
        # revise-first：尽量用全版一次改到位；balanced：低/中用简版，高用全版
        if policy == "revise-first":
            new = full or simple
        else:
            new = full if risk == "高" else simple
        if not new:
            new = simple or full
        item["action"] = "replace"
        item["old"] = old
        item["new"] = new
        plan.append(item)

    return plan


# ============================================================
# OOXML 精确 in-place 修改
# ============================================================

def _run_text(run_elem):
    """取 w:r 的文本（来自其直接子 <w:t>）"""
    return "".join(t.text or "" for t in run_elem.findall(f"{{{W_NS}}}t"))


def _set_run_text(run_elem, text):
    """清空 run 内所有 <w:t> 并写入新文本（保留 rPr 格式）"""
    for t in run_elem.findall(f"{{{W_NS}}}t"):
        run_elem.remove(t)
    if not text:
        return
    t = OxmlElement("w:t")
    t.set(f"{{{XML_SPACE}}}space", "preserve")
    t.text = text
    rpr = run_elem.find(f"{{{W_NS}}}rPr")
    if rpr is not None:
        rpr.addnext(t)
    else:
        run_elem.insert(0, t)


def _make_del_run(text, src_rpr, author, date, wid):
    """构造 <w:del> 删除标记（含内部 <w:r><w:delText>）"""
    del_ = OxmlElement("w:del")
    del_.set(f"{{{W_NS}}}id", str(wid))
    del_.set(f"{{{W_NS}}}author", author)
    del_.set(f"{{{W_NS}}}date", date)
    r = OxmlElement("w:r")
    if src_rpr is not None:
        r.append(copy.deepcopy(src_rpr))
    dt = OxmlElement("w:delText")
    dt.set(f"{{{XML_SPACE}}}space", "preserve")
    dt.text = text
    r.append(dt)
    del_.append(r)
    return del_


def _make_ins_run(text, src_rpr, author, date, wid):
    """构造 <w:ins> 插入标记（含内部 <w:r><w:t>）"""
    ins = OxmlElement("w:ins")
    ins.set(f"{{{W_NS}}}id", str(wid))
    ins.set(f"{{{W_NS}}}author", author)
    ins.set(f"{{{W_NS}}}date", date)
    r = OxmlElement("w:r")
    if src_rpr is not None:
        r.append(copy.deepcopy(src_rpr))
    t = OxmlElement("w:t")
    t.set(f"{{{XML_SPACE}}}space", "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    return ins


def _para_full_text(p):
    """取段落全文（含 ins/del 内文本，便于定位）"""
    elem = p._p
    parts = []
    for t in elem.iter(f"{{{W_NS}}}t"):
        parts.append(t.text or "")
    for t in elem.iter(f"{{{W_NS}}}delText"):
        parts.append(t.text or "")
    return "".join(parts)


def _replace_in_paragraph(para_elem, old, new, author, date, wid):
    """
    在段落内精确替换 old -> <w:del>old</w:del> + <w:ins>new</w:ins>。
    支持 old 跨多个 run；new 为空时仅生成 del（删除）。
    返回是否成功定位并替换。
    """
    if not old:
        return False

    runs = para_elem.findall(f"{{{W_NS}}}r")
    if not runs:
        return False

    full = ""
    run_texts = []
    for r in runs:
        rt = _run_text(r)
        run_texts.append(rt)
        full += rt

    idx = full.find(old)
    if idx == -1:
        return False

    # 定位 start/end run（含字符偏移）
    pos = 0
    start_ri = end_ri = -1
    start_off = end_off = 0
    for i, rt in enumerate(run_texts):
        if start_ri == -1 and pos <= idx < pos + len(rt):
            start_ri = i
            start_off = idx - pos
        if pos < idx + len(old) <= pos + len(rt):
            end_ri = i
            end_off = (idx + len(old)) - pos
            break
        pos += len(rt)
    if start_ri == -1 or end_ri == -1:
        return False

    src_rpr = runs[start_ri].find(f"{{{W_NS}}}rPr")

    # ---------- 单 run 情况 ----------
    if start_ri == end_ri:
        rt = run_texts[start_ri]
        before = rt[:start_off]
        after = rt[end_off:]

        _set_run_text(runs[start_ri], before)  # start run 现在只保留 before 部分

        del_ = _make_del_run(old, src_rpr, author, date, wid) if old else None
        if del_ is not None:
            runs[start_ri].addnext(del_)

        if new:
            ins_ = _make_ins_run(new, src_rpr, author, date, wid)
            if del_ is not None:
                del_.addnext(ins_)
            else:
                runs[start_ri].addnext(ins_)

        if after:
            after_run = OxmlElement("w:r")
            if src_rpr is not None:
                after_run.append(copy.deepcopy(src_rpr))
            _set_run_text(after_run, after)
            if new:
                ins_.addnext(after_run)
            elif del_ is not None:
                del_.addnext(after_run)
            else:
                runs[start_ri].addnext(after_run)

        if not before:
            para_elem.remove(runs[start_ri])
        return True

    # ---------- 多 run 情况 ----------
    _set_run_text(runs[start_ri], run_texts[start_ri][:start_off])
    _set_run_text(runs[end_ri], run_texts[end_ri][end_off:])

    del_ = _make_del_run(old, src_rpr, author, date, wid) if old else None
    if del_ is not None:
        runs[start_ri].addnext(del_)
        if new:
            ins_ = _make_ins_run(new, src_rpr, author, date, wid)
            del_.addnext(ins_)

    # 删除中间 run
    for i in range(start_ri + 1, end_ri):
        para_elem.remove(runs[i])

    if not run_texts[start_ri][:start_off]:
        para_elem.remove(runs[start_ri])
    if not run_texts[end_ri][end_off:]:
        para_elem.remove(runs[end_ri])
    return True


def _insert_after_keyword(doc, keyword, new, author, date, wid):
    """在含 keyword 的段落之后插入一个新段落（内容用 <w:ins> 包裹）"""
    target = None
    for p in doc.paragraphs:
        if keyword and keyword in _para_full_text(p):
            target = p
            break
    new_p = doc.add_paragraph()  # 先添加到末尾
    new_elem = new_p._p
    if target is not None:
        target._p.addnext(new_elem)
    # 清空 python-docx 可能生成的空 run，写入 ins
    for r in new_elem.findall(f"{{{W_NS}}}r"):
        new_elem.remove(r)
    new_elem.append(_make_ins_run(new, None, author, date, wid))
    return True


# ============================================================
# 应用 render plan 到 docx
# ============================================================

def apply_plan(doc, plan, author, date):
    """按 render plan 确定性应用到 docx，返回执行结果列表"""
    results = []
    wid = 1

    for item in plan:
        action = item.get("action")
        rid = item.get("id")

        if action == "comment":
            results.append({
                "id": rid, "action": "comment", "applied": False,
                "note": "仅进意见书，不落痕",
            })
            continue

        if action in ("replace", "delete"):
            old = item.get("old", "")
            new = item.get("new", "") if action == "replace" else ""
            found = False
            for p in doc.paragraphs:
                if old and old in _para_full_text(p):
                    ok = _replace_in_paragraph(p._p, old, new, author, date, wid)
                    if ok:
                        found = True
                        wid += 1
                        break
            results.append({
                "id": rid, "action": action, "applied": found,
                "note": "" if found else "原文未定位（replace 目标文本未命中）",
            })

        elif action == "insert":
            new = item.get("new", "")
            kw = item.get("anchor", "")
            ok = _insert_after_keyword(doc, kw, new, author, date, wid)
            results.append({
                "id": rid, "action": "insert", "applied": ok,
                "note": "" if ok else "未找到插入锚点",
            })
            if ok:
                wid += 1

    return results


# ============================================================
# 归档
# ============================================================

def archive(plan, results, out_path, plan_out, doc_name):
    """写出 render plan + 执行日志，便于复核与回放"""
    log = {
        "generated_at": datetime.datetime.now().isoformat(timespec="seconds"),
        "source_doc": doc_name,
        "output_doc": os.path.basename(out_path),
        "plan_summary": {
            "total": len(plan),
            "revise": sum(1 for i in plan if i["action"] in ("replace", "delete", "insert")),
            "comment": sum(1 for i in plan if i["action"] == "comment"),
        },
        "execution": results,
    }
    if plan_out:
        with open(plan_out, "w", encoding="utf-8") as f:
            json.dump(plan, f, ensure_ascii=False, indent=2)
    log_path = out_path + ".exec.log.json"
    with open(log_path, "w", encoding="utf-8") as f:
        json.dump(log, f, ensure_ascii=False, indent=2)
    return log_path


# ============================================================
# 主入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="台账驱动的确定性修订稿生成（hybrid）")
    parser.add_argument("--ledger", required=True, help="review-issues-contract.json 路径")
    parser.add_argument("--doc", required=True, help="原始合同 docx 路径")
    parser.add_argument("--out", required=True, help="输出修订稿 docx 路径")
    parser.add_argument("--policy", default="balanced",
                        choices=["revise-first", "balanced", "comment-first"],
                        help="edit_policy（默认 balanced）")
    parser.add_argument("--plan-out", default=None, help="render plan 输出 JSON 路径（可选）")
    parser.add_argument("--author", default="合同智囊·WorkBuddy", help="修订作者署名")
    args = parser.parse_args()

    date = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    if not os.path.exists(args.ledger):
        print(f"❌ 账本不存在: {args.ledger}")
        sys.exit(1)
    if not os.path.exists(args.doc):
        print(f"❌ 原文件不存在: {args.doc}")
        sys.exit(1)

    ledger = load_ledger(args.ledger)
    plan = build_render_plan(ledger, args.policy)
    if not plan:
        print("⚠️ 账本中无 status=confirmed 的问题，未生成修订。")
        sys.exit(0)

    doc = Document(args.doc)
    results = apply_plan(doc, plan, args.author, date)
    doc.save(args.out)

    log_path = archive(plan, results, args.out, args.plan_out,
                       os.path.basename(args.doc))

    applied = sum(1 for r in results if r["applied"])
    commented = sum(1 for r in results if r["action"] == "comment")
    missed = [r for r in results if r["action"] != "comment" and not r["applied"]]

    print(f"✅ 修订稿已生成: {args.out}")
    print(f"   render plan 条数: {len(plan)}（落痕 {applied} / 仅批注 {commented}）")
    if missed:
        print(f"⚠️ 以下条目原文未定位，未落痕：")
        for m in missed:
            print(f"   - {m['id']} ({m['action']}): {m['note']}")
    print(f"📦 执行日志: {log_path}")
    if args.plan_out:
        print(f"📦 render plan: {args.plan_out}")


if __name__ == "__main__":
    main()
